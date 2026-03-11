
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create new document
doc = Document()

# Title
title = doc.add_heading('Business Summary Report: Predictive Insights for Collections Strategy', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Summary of Predictive Insights
doc.add_heading('1. Summary of Predictive Insights', level=1)
doc.add_paragraph('Our predictive modeling analysis has identified key behavioral and demographic indicators that signal a high probability of future delinquency. By leveraging these insights, the Collections team can shift from reactive chasing to proactive prevention.')

# Bullet points
p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run('Top Predictor: ')
run.bold = True
p.add_run('Credit Utilization is the leading risk indicator. Customers crossing the 50% utilization threshold show a significantly accelerated risk of default.')

p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run('High-Risk Segment: ')
run.bold = True
p.add_run('The "Young Adult" segment (<30 years old) exhibits a delinquency rate of 20.8%, nearly double that of middle-aged customers (11.1%).')

p = doc.add_paragraph()
p.style = 'List Bullet'
run = p.add_run('Hidden Signal: ')
run.bold = True
p.add_run('Payment behavior from 4-5 months ago is more predictive of current default than the most recent month, suggesting a "lag effect" in financial stress.')

# Key Insights Table
doc.add_heading('Key Insights Summary Table', level=2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Key Insight'
hdr_cells[1].text = 'Customer Segment'
hdr_cells[2].text = 'Influencing Variables'
hdr_cells[3].text = 'Potential Business Impact'

# Make header bold
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

rows = [
    ('Utilization Threshold', 'Customers >50% Utilization', 'Credit Utilization, Debt-to-Income', 'Early intervention opportunity before default occurs.'),
    ('Income Vulnerability', 'Low Income (<$40k)', 'Income, Employment Status', 'Need for affordable repayment plans vs. aggressive collection.'),
    ('Card Type Risk', 'Business Card Holders', 'Credit Card Type', 'Targeted usage reviews for business accounts.')
]

for insight, segment, vars, impact in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = insight
    row_cells[1].text = segment
    row_cells[2].text = vars
    row_cells[3].text = impact

# 2. Recommendation Framework
doc.add_heading('2. Recommendation Framework', level=1)
doc.add_paragraph('Based on the verified link between high utilization and default, we propose a proactive engagement strategy.')

# Subheadings
h = doc.add_heading('Restated Insight:', level=3)
doc.add_paragraph('High credit utilization (>50%) is the strongest behavioral predictor of pending delinquency.')

h = doc.add_heading('Proposed Recommendation:', level=3)
doc.add_paragraph('Pilot a "Credit Health" automated SMS alert system targeting customers who exceed 45% credit utilization.')

# SMART Goals
doc.add_heading('SMART Analysis:', level=3)
p = doc.add_paragraph()
p.add_run('Specific: ').bold = True
p.add_run('Target active customers with utilization between 45-60%.')

p = doc.add_paragraph()
p.add_run('Measurable: ').bold = True
p.add_run('Aim to reduce the roll-rate to 30+ day delinquency by 10% in the pilot group.')

p = doc.add_paragraph()
p.add_run('Actionable: ').bold = True
p.add_run('Utilize existing CRM marketing automation tools to trigger SMS based on daily balance updates.')

p = doc.add_paragraph()
p.add_run('Relevant: ').bold = True
p.add_run('Directly addresses the #1 predictive risk factor found in the model.')

p = doc.add_paragraph()
p.add_run('Time-bound: ').bold = True
p.add_run('Launch 3-month pilot starting Q2, with monthly performance reviews.')

h = doc.add_heading('Justification and Business Rationale:', level=3)
doc.add_paragraph('This approach aligns with Geldium’s goal of reducing risk while maintaining customer relationships. Unlike credit line cuts, which can cause frustration, data-driven education empowers customers to self-correct. It is low-cost (<$0.01 per message) with a high potential ROI through avoided charge-offs.')

# 3. Ethical and Responsible AI Considerations
doc.add_heading('3. Ethical and Responsible AI Considerations', level=1)

p = doc.add_paragraph()
p.add_run('Fairness & Bias Mitigations: ').bold = True
p.add_run('Our analysis detected that younger adults and lower-income groups have higher predicted risk scores. To prevent systematic denial of service, we recommend:')
p.style = 'List Paragraph' 
# sub-bullets
p = doc.add_paragraph('• Regularly auditing model performance (Recall/False Positive Rate) across demographic groups.')
p.style = 'List Bullet'
p = doc.add_paragraph('• Excluding "Location" from the final model to avoid geographic redlining, as Los Angeles residents were disproportionately flagged.')
p.style = 'List Bullet'

p = doc.add_paragraph()
p.add_run('Explainability & Transparency: ').bold = True
p.add_run('We utilized SHAP (SHapley Additive exPlanations) values to ensure the model is not a "black box." Every risk score can be decomposed into its contributing factors, allowing customer service agents to explain clearly why an account is flagged (e.g., "Your high utilization is the main factor").')

p = doc.add_paragraph()
p.add_run('Responsible Financial Decision-Making: ').bold = True
p.add_run('The recommendation emphasizes "Nudges" (education) over "Shoves" (punitive action). This supports responsible usage habits rather than creating a debt trap.')

# Save
file_name = 'Final_Business_Summary_Report.docx'
doc.save(file_name)
print(f"Successfully generated {file_name}")
