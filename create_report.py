
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Create a new Document
    doc = Document()

    # Title
    title = doc.add_heading('Exploratory Data Analysis (EDA) Summary Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta info
    p = doc.add_paragraph()
    p.add_run('Project: ').bold = True
    p.add_run('Delinquency Prediction Model\n')
    p.add_run('Prepared For: ').bold = True
    p.add_run('Tata iQ Analytics Team & Geldium Decision-Makers\n')
    p.add_run('Date: ').bold = True
    p.add_run('October 26, 2023')
    
    doc.add_heading('1. Dataset Review and Key Insights', level=1)
    
    p = doc.add_paragraph('The initial inspection of the ')
    p.add_run('Delinquency_prediction_dataset').italic = True
    p.add_run(' revealed technical corruption in the file header, requiring a manual recovery of the data. After successfully extracting 500 records, the following quality issues and insights were identified:')

    # List of issues
    doc.add_heading('Data Quality Issues:', level=2)
    
    items = [
        ("Inconsistent Categorical Data", 'The "Employment_Status" column contains duplicate categories due to inconsistent naming conventions (e.g., "Employed", "employed", "EMP"). These must be standardized before modeling.'),
        ("Counter-Intuitive Correlations", "Preliminary analysis highlights potential data validity issues. For instance, 'Missed_Payments' shows a slight negative correlation (-0.026) with 'Delinquent_Account', which contradicts the expected behavioral pattern."),
        ("Weak Predictors", "Most numeric variables (Income, Credit Score) show very weak correlations (<0.05) with the target variable, suggesting non-linear relationships or that the dataset may need additional feature engineering.")
    ]
    
    for title, text in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title + ": ").bold = True
        p.add_run(text)

    # Risk Indicators
    doc.add_heading('Early Risk Indicators:', level=2)
    
    risks = [
        ("Credit Utilization", "Customers with delinquent accounts have a slightly higher average Credit_Utilization (50.6%) compared to non-delinquent customers (48.8%)."),
        ("Income", "Surprisingly, there is a weak positive correlation between Income and Delinquency, warranting further segmentation analysis.")
    ]
    
    for title, text in risks:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title + ": ").bold = True
        p.add_run(text)

    doc.add_heading('2. Missing Data & Imputation Strategy', level=1)
    
    doc.add_paragraph('The analysis identified missing values in three critical columns. Below is the proposed strategy to handle them.')

    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Missing Count'
    hdr_cells[2].text = 'Handling Method'
    hdr_cells[3].text = 'Justification'
    
    # Header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    data = [
        ('Income', '39 (7.8%)', 'Median Imputation', 'Income data is skewed with outliers. Median provides a robust central tendency.'),
        ('Loan_Balance', '29 (5.8%)', 'Median Imputation', 'Similar to income, balances vary significantly. Median preserves the customer profile.'),
        ('Credit_Score', '2 (0.4%)', 'Mean Imputation', 'With very few missing values and normal distribution, mean is a safe estimate.')
    ]

    for v, c, m, j in data:
        row_cells = table.add_row().cells
        row_cells[0].text = v
        row_cells[1].text = c
        row_cells[2].text = m
        row_cells[3].text = j

    doc.add_heading('3. Pattern Detection & Risk Factors', level=1)
    
    doc.add_heading('High-Risk Indicators:', level=2)
    
    high_risks = [
        ("Elevated Credit Utilization", "A utilization rate above 50% is a consistent marker for delinquency."),
        ("Employment Status 'Unemployed'", "This segment (~18.6%) historically carries higher risk inherent to income instability."),
        ("Low Account Tenure", "Newer accounts often carry higher risk. The negative correlation (-0.039) suggests newer customers are slightly more prone to delinquency.")
    ]
    
    for title, text in high_risks:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title + ": ").bold = True
        p.add_run(text)
        
    doc.add_heading('Unexpected Findings for Investigation:', level=2)
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Negative Correlation of Missed Payments: ").bold = True
    p.add_run("This implies customers with fewer missed payments were more likely to be delinquent. This anomaly requires immediate validation with the data engineering team.")

    doc.add_heading('4. Conclusion & Next Steps', level=1)
    
    steps = [
        "Standardize existing categorical labels (Employment Status).",
        "Impute missing financial data using the Median strategy.",
        "Investigate the 'Missed Payments' variable correlations.",
        "Feature Engineering: Create new features to capture risk better."
    ]
    
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    # Save
    doc.save('EDA_Summary_Report_Final.docx')
    print("Report generated successfully: EDA_Summary_Report_Final.docx")

except ImportError:
    print("python-docx not found. Creating a rich text HTML file instead.")
    
    html_content = """
    <html>
    <body style="font-family: Arial, sans-serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px;">
        <h1 style="text-align: center;">Exploratory Data Analysis (EDA) Summary Report</h1>
        <p><strong>Project:</strong> Delinquency Prediction Model<br>
        <strong>Prepared For:</strong> Tata iQ Analytics Team & Geldium Decision-Makers<br>
        <strong>Date:</strong> October 26, 2023</p>
        
        <hr>
        
        <h2>1. Dataset Review and Key Insights</h2>
        <p>The initial inspection of the <em>Delinquency_prediction_dataset</em> revealed technical corruption in the file header, requiring a manual recovery of the data. After successfully extracting 500 records, the following quality issues and insights were identified:</p>
        
        <h3>Data Quality Issues:</h3>
        <ul>
            <li><strong>Inconsistent Categorical Data:</strong> The "Employment_Status" column contains duplicate categories due to inconsistent naming conventions (e.g., "Employed", "employed", "EMP"). These must be standardized before modeling.</li>
            <li><strong>Counter-Intuitive Correlations:</strong> Preliminary analysis highlights potential data validity issues. For instance, "Missed_Payments" shows a slight negative correlation (-0.026) with "Delinquent_Account", which contradicts the expected behavioral pattern.</li>
            <li><strong>Weak Predictors:</strong> Most numeric variables (Income, Credit Score) show very weak correlations (<0.05) with the target variable.</li>
        </ul>

        <h2>2. Missing Data & Imputation Strategy</h2>
        <table border="1" cellpadding="10" style="border-collapse: collapse; width: 100%;">
            <tr style="background-color: #f2f2f2;">
                <th>Variable</th>
                <th>Missing Count</th>
                <th>Handling Method</th>
                <th>Justification</th>
            </tr>
            <tr>
                <td>Income</td>
                <td>39 (7.8%)</td>
                <td>Median Imputation</td>
                <td>Income data is skewed with outliers.</td>
            </tr>
            <tr>
                <td>Loan_Balance</td>
                <td>29 (5.8%)</td>
                <td>Median Imputation</td>
                <td>Median preserves the typical customer profile better than mean.</td>
            </tr>
            <tr>
                <td>Credit_Score</td>
                <td>2 (0.4%)</td>
                <td>Mean Imputation</td>
                <td>Normal distribution allows for mean imputation.</td>
            </tr>
        </table>

        <h2>3. Pattern Detection & Risk Factors</h2>
        <h3>High-Risk Indicators:</h3>
        <ul>
            <li><strong>Elevated Credit Utilization:</strong> A utilization rate above 50% is a consistent marker for delinquency.</li>
            <li><strong>Employment Status "Unemployed":</strong> Represents ~18.6% of the dataset and carries inherent income instability risks.</li>
            <li><strong>Low Account Tenure:</strong> Newer accounts often carry higher risk.</li>
        </ul>
        
        <h2>4. Conclusion & Next Steps</h2>
        <ol>
            <li>Standardize existing categorical labels.</li>
            <li>Impute missing financial data using the Median strategy.</li>
            <li>Investigate the "Missed Payments" variable.</li>
            <li>Perform Feature Engineering.</li>
        </ol>
    </body>
    </html>
    """
    
    with open('EDA_Summary_Report_Final.html', 'w') as f:
        f.write(html_content)
    print("Report generated successfully: EDA_Summary_Report_Final.html (Open in browser/Word)")
