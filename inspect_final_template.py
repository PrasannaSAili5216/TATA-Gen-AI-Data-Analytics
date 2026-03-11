from docx import Document

try:
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    doc = Document('TASK_DETAILS/TASK_3/Updated_Business_Summary_Report_Template.docx')
    print("--- Template Structure ---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"{i}: {text}")
    
    print("\n--- Tables ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for row in table.rows:
            # Join cell text with pipe for clearer table structure view
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            print(f"  {row_text}")
            
except Exception as e:
    print(f"Error reading template: {e}")
