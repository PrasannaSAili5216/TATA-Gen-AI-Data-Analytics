
import sys

# Redirect stdout to a file with utf-8 encoding
sys.stdout = open('template_structure_utf8.txt', 'w', encoding='utf-8')

try:
    from docx import Document
    
    doc = Document('EDA_SummaryReport_Template.docx')
    
    print("--- Template Structure ---")
    for para in doc.paragraphs:
        if para.text.strip():
            print(f"Style: {para.style.name} | Text: {para.text}")
            
    print("\n--- Tables ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i+1} rows: {len(table.rows)}")
        row1 = [cell.text for cell in table.rows[0].cells]
        print(f"Header: {row1}")

except Exception as e:
    print(f"Error reading template: {e}")
