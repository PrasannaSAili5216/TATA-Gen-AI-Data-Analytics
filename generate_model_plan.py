
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_model_plan():
    doc = Document()
    
    # Title
    heading = doc.add_heading('Predictive Model Plan – Delinquency Prediction', level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Model Logic
    doc.add_heading('1. Model Logic (Generated with GenAI)', level=1)
    doc.add_paragraph('Below is the step-by-step logic and pseudo-code for the proposed Random Forest prediction model.')
    
    logic_code = """
# Pseudo-Code for Delinquency Prediction Model

1. Data Loading & Preprocessing
   - Load 'cleaned_dataset.csv' (output from Task 1)
   - Encode Categorical Variables:
     - One-Hot Encode: 'Employment_Status', 'Credit_Card_Type', 'Location'
   - Scale Numerical Variables: 'Income', 'Loan_Balance', 'Credit_Utilization'

2. Feature Engineering
   - Create 'Debt_to_Income_Ratio' = Loan_Balance / Income
   - Select Features (X) and Target (y = 'Delinquent_Account')

3. Train-Test Split
   - Split data into Training (80%) and Testing (20%) sets
   - Stratify split to maintain delinquency proportion

4. Handling Class Imbalance (Crucial for Fraud/Delinquency)
   - Apply SMOTE (Synthetic Minority Over-sampling Technique) to Training Data ONLY
   - Goal: Balance the count of Delinquent vs. Non-Delinquent cases in training

5. Model Initialization & Training
   - Algorithm: Random Forest Classifier
   - Hyperparameters: n_estimators=100, max_depth=10, random_state=42
   - Train: model.fit(X_train_resampled, y_train_resampled)

6. Prediction & Explanation
   - Generate Predictions: y_pred = model.predict(X_test)
   - Probabilities: y_prob = model.predict_proba(X_test)
   - Explainability: Calculate Feature Importance scores
"""
    p = doc.add_paragraph(logic_code)
    p.style = 'No Spacing'
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # 2. Justification for Model Choice
    doc.add_heading('2. Justification for Model Choice', level=1)
    
    doc.add_paragraph('Selected Model: Random Forest Classifier', style='List Bullet')
    
    p = doc.add_paragraph()
    p.add_run('Accuracy & Robustness: ').bold = True
    p.add_run('Random Forest is an ensemble method that aggregates multiple decision trees. This makes it highly accurate and resistant to overfitting compared to single decision trees. It handles non-linear relationships well, such as the non-linear impact of age or income on default risk.')
    
    p = doc.add_paragraph()
    p.add_run('Transparency & Explainability: ').bold = True
    p.add_run('Unlike "black box" models (e.g., Deep Neural Networks), Random Forest provides "Feature Importance" metrics. This allows us to explain exactly which factors (e.g., Credit Utilization, Income) deal the most damage to a credit score, which is a regulatory requirement in lending (Right to Explanation).')
    
    p = doc.add_paragraph()
    p.add_run('Suitability for Geldium: ').bold = True
    p.add_run('The dataset contains both categorical (Employment) and numerical data. Random Forest handles mixed data types effectively without extensive normalization. It is also robust to outliers, which were observed in the Income variable during EDA.')

    # 3. Evaluation Strategy
    doc.add_heading('3. Evaluation Strategy', level=1)
    
    doc.add_paragraph('Metrics for Success:', style='Heading 2')
    
    p = doc.add_paragraph()
    p.add_run('Recall (Sensitivity): ').bold = True
    p.add_run('This is the PRIMARY metric. In delinquency prediction, a False Negative (missed delinquent customer) is widely more expensive than a False Positive (investigating a safe customer). We aim to maximize Recall to catch at least 85% of potential defaults.')
    
    p = doc.add_paragraph()
    p.add_run('ROC-AUC Score: ').bold = True
    p.add_run('Used to evaluate the model’s ability to distinguish between high-risk and low-risk customers across different probability thresholds.')
    
    # 4. Ethical Governance & Compliance Framework
    doc.add_heading('4. Ethical Governance & Compliance Framework', level=1)
    doc.add_paragraph('To ensure the system operates responsibly and legally, we implement a framework aligned with key financial regulations.')

    # Regulatory Alignment Table
    doc.add_heading('Regulatory Alignment Strategy', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Regulation'
    hdr[1].text = 'Focus Area'
    hdr[2].text = 'System Implementation'
    
    reg_data = [
        ('ECOA (US)', 'Anti-Discrimination', 'Automated nightly bias testing (Disparate Impact Analysis) across protected groups.'),
        ('GDPR (EU/UK)', 'Right to Explanation', 'SHAP-based "Reason Codes" provided for every adverse decision.'),
        ('FCA (UK)', 'Treating Customers Fairly', 'Model prioritization of "Support" interventions (forbearance) over "Collections" for vulnerable customers.'),
        ('FCRA (US)', 'Data Accuracy', 'Real-time data ingestion to ensure decisions use the latest credit file; "Dispute" mechanisms act as a circuit breaker.')
    ]
    
    for reg, focus, impl in reg_data:
        row = table.add_row().cells
        row[0].text = reg
        row[1].text = focus
        row[2].text = impl

    # Governance Pillars
    doc.add_heading('Governance Pillars', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Human-in-the-Loop: ').bold = True
    p.add_run('The AI is an advisor. Adverse actions (e.g., closures) REQUIRE human approval.')

    p = doc.add_paragraph()
    p.add_run('Audit Trails: ').bold = True
    p.add_run('Immutable "Decision Ledger" logs every input, risk score, and action for regulatory review.')
    
    p = doc.add_paragraph()
    p.add_run('Customer Recourse: ').bold = True
    p.add_run('Clear appeals process for customers to challenge AI-driven decisions with new info.')

    # Practical Compliance Implementation (Lifecycle Strategy)
    doc.add_heading('Key Implementation Strategies', level=2)
    doc.add_paragraph('To embed compliance into the daily operation of the AI system, we adopt the following lifecycle strategies:')

    strategies = [
        ('Map the Decision Flow', 'Visual documentation of how data inputs (credit reports) flow to outputs (risk flags), ensuring every decision path is traceable and explainable in simple terms.'),
        ('Automated Compliance Checks', 'The pipeline includes "Circuit Breakers". If the model flags a disproportionate number of protected class customers in a single batch, execution halts for manual review.'),
        ('Early Compliance Engagement', 'Legal and Risk teams are involved in the "Feature Selection" phase to veto risky variables (e.g., Zip Code) before training begins.'),
        ('Maintain Documentation', 'Detailed records of training data, feature selection, evaluation metrics, and model versions are maintained to support internal and external regulatory reviews.'),
        ('Monitor & Adapt', 'Regular audits assess whether real-world conditions (e.g., economic downturn) have introduced new biases, ensuring the model evolves with the regulatory landscape.')
    ]

    for title, desc in strategies:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{title}: ').bold = True
        p.add_run(desc)

    # 5. Operational Implementation Strategy
    doc.add_heading('5. Operational Implementation Strategy', level=1)
    
    doc.add_paragraph('To operationalize the model, we will deploy an Agentic AI Debt-Management System. Unlike static rule-based automation, this system possesses the autonomy to reason, adapt, and balance competing business objectives (e.g., risk mitigation vs. customer retention).')

    # Component 1: Data Pipeline & Decision Engine
    doc.add_heading('Component 1: agentic Decision Engine', level=2)
    p = doc.add_paragraph()
    p.add_run('Contextual Reasoning: ').bold = True
    p.add_run('The Agentic AI ingests real-time data (transactions, credit bureau) to build a dynamic customer profile. It looks beyond simple risk scores to understand *context*. For example, if a "High Risk" customer shows a sudden drop in income, the Agent infers financial distress and prioritizes "Support" over "Collections".')
    
    p = doc.add_paragraph()
    p.add_run('Balancing Objectives: ').bold = True
    p.add_run('The engine weighs multiple factors simultaneously to determine the optimal action that maximizes recovery while minimizing customer churn.')

    # Component 2: Action Layer (Interventions)
    doc.add_heading('Component 2: Autonomous Action Layer', level=2)
    doc.add_paragraph('The system autonomously executes personalized interventions, adjusting the channel and tone based on predicted receptiveness:')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Low Risk (<20%): ').bold = True
    p.add_run('Monitor Only. The agent decides to withhold action to avoid unnecessary friction.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Medium Risk (20-60%): ').bold = True
    p.add_run('Dynamic Nudge. The agent selects the optimal channel (SMS/Email) and time (e.g., 6 PM). If an email is ignored, it autonomously switches strategies to a polite SMS reminder.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('High Risk (>60%): ').bold = True
    p.add_run('Intervention or Escalation. For confirmed distress, the Agent autonomously offers a "Payment Deferral". For unexplained non-payment, it drafts a case for the Collections Team.')

    # Component 3: Learning Loop
    doc.add_heading('Component 3: Continuous Learning Loop', level=2)
    p = doc.add_paragraph()
    p.add_run('Reinforcement Learning: ').bold = True
    p.add_run('Outcomes serve as feedback. If "Offer Deferral" leads to higher long-term recovery than "Demand Payment", the Agent reinforces this pathway for similar future profiles.')
    
    p = doc.add_paragraph()
    p.add_run('Adaptive Retraining: ').bold = True
    p.add_run('The core Random Forest model is retrained quarterly, but the Agent’s decision logic adapts daily based on immediate campaign results.')

    # Oversight & Guardrails
    doc.add_heading('Human Oversight & Ethical Guardrails', level=2)
    p = doc.add_paragraph()
    p.add_run('Human-in-the-Loop: ').bold = True
    p.add_run('The Agent is autonomous but accountable. Adverse actions (e.g., reporting to bureaus) require human approval. Complex edge cases are routed to human specialists.')
    
    p = doc.add_paragraph()
    p.add_run('Compliance Constraints: ').bold = True
    p.add_run('Hard-coded "Do Not Cross" lines ensure the Agent never violates regulations (e.g., calling outside legal hours or using aggressive language).')

    # Save
    start_path = 'Task_2_Model_Plan_Submission.docx'
    doc.save(start_path)
    print(f"Document saved as {start_path}")

if __name__ == "__main__":
    create_model_plan()
