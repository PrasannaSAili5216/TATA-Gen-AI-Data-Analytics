
from docx import Document
import sys
import os

def read_docx(file_path, output_path):
    if not os.path.exists(file_path):
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"Error: File not found at {file_path}")
        return

    try:
        doc = Document(file_path)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"--- Content of {os.path.basename(file_path)} ---\n")
            for para in doc.paragraphs:
                cleaned_text = para.text.strip()
                if cleaned_text:
                    f.write(cleaned_text + "\n")
            
            f.write("\n--- Tables ---\n")
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    f.write(" | ".join(row_text) + "\n")
                f.write("-" * 20 + "\n")
            
        print(f"Successfully wrote to {output_path}")
            
    except Exception as e:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"Error reading docx: {e}")

if __name__ == "__main__":
    read_docx(r"d:\Projects\Forage TATA Intern\TEMPLATE_FILES\TASK_2\Task 2_ModelPlan_Template.docx", "template_content_utf8.txt")
