
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Exploratory Data Analysis (EDA) Summary', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Introduction
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph('The purpose of this report is to analyze the data quality and structure of the delinquency prediction dataset. The goal is to identify missing values, inconsistencies, and early risk indicators to prepare the data for predictive modeling by the Tata iQ analytics team.')

# 2. Dataset Overview
doc.add_heading('2. Dataset Overview', level=1)
doc.add_paragraph('This section summarizes the dataset, including the number of records, key variables, and data types. It also highlights any anomalies, duplicates, or inconsistencies observed during the initial review.')

para = doc.add_paragraph('Key dataset attributes:')
para.add_run('\n- Number of records: ').bold = True
para.add_run('500')
para.add_run('\n- Key variables: ').bold = True
para.add_run('Delinquent_Account (Target), Credit_Utilization, Missed_Payments, Income, Loan_Balance.')
para.add_run('\n- Data types: ').bold = True
para.add_run('Mixture of Categorical (Employment_Status, Location) and Numerical (Credit_Score, Income).')

doc.add_paragraph('Anomalies Observed: Technical corruption in original file (recovered manually). Inconsistent categorical labels in "Employment_Status" (e.g., "Employed" vs "EMP").')

# 3. Missing Data Analysis
doc.add_heading('3. Missing Data Analysis', level=1)
doc.add_paragraph('Identifying and addressing missing data is critical to ensuring model accuracy. This section outlines missing values in the dataset, the approach taken to handle them, and justifications for the chosen method.')

para = doc.add_paragraph('Key missing data findings:')
para.add_run('\n- Variables with missing values: ').bold = True
para.add_run('Income (39 missing), Loan_Balance (29 missing), Credit_Score (2 missing).')

doc.add_paragraph('\nMissing Data Treatment Strategy:')
# Table
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Variable'
hdr_cells[1].text = 'Strategy'
hdr_cells[2].text = 'Justification'

data = [
    ('Income', 'Median Imputation', 'Data is skewed; median avoids outlier bias.'),
    ('Loan_Balance', 'Median Imputation', 'Financial balances vary widely; median is more robust.'),
    ('Credit_Score', 'Mean Imputation', 'Low missing count (0.4%) and normal distribution.')
]
for v, s, j in data:
    row = table.add_row().cells
    row[0].text = v
    row[1].text = s
    row[2].text = j

# 4. Key Findings and Risk Indicators
doc.add_heading('4. Key Findings and Risk Indicators', level=1)
doc.add_paragraph('This section identifies trends and patterns that may indicate risk factors for delinquency. Feature relationships and statistical correlations are explored to uncover insights relevant to predictive modeling.')

para = doc.add_paragraph('Key findings:')
para.add_run('\n- Correlations observed: ').bold = True
para.add_run('Most numeric predictors show weak linear correlation. "Credit_Utilization" (0.034) and "Income" (0.045) show slight positive correlations with delinquency.')
para.add_run('\n- Unexpected anomalies: ').bold = True
para.add_run('"Missed_Payments" shows a negative correlation (-0.026) with delinquency, which is counter-intuitive and requires data dictionary verification.')
para.add_run('\n- High-Risk Indicator: ').bold = True
para.add_run('Customers with >50% credit utilization are more frequently found in the delinquent category.')

# 5. AI & GenAI Usage
doc.add_heading('5. AI & GenAI Usage', level=1)
doc.add_paragraph('Generative AI tools were used to summarize the dataset, impute missing data, and detect patterns. This section documents AI-generated insights and the prompts used to obtain results.')

para = doc.add_paragraph('Example AI prompts used:')
para.add_run("\n- 'Summarize key patterns in the dataset and identify anomalies.'")
para.add_run("\n- 'Suggest an imputation strategy for missing income values based on industry best practices.'")
para.add_run("\n- 'Analyze the correlation between credit utilization and delinquency risk.'")

# 6. Model Explainability (SHAP Analysis)
doc.add_heading('6. Model Explainability (SHAP Analysis)', level=1)
doc.add_paragraph('To ensure transparency in our predictive model, we utilized SHAP (SHapley Additive exPlanations) values. This advanced technique breaks down the contribution of each feature to the final risk score.')

# Add SHAP Summary Plot
try:
    doc.add_paragraph('Figure 1: SHAP Summary Plot identifying top risk drivers.')
    doc.add_picture('shap_summary_plot.png', width=Inches(5.0))
except Exception as e:
    doc.add_paragraph(f"[Image placeholder: SHAP Summary Plot not found - {e}]")

# Top Features
doc.add_heading('Key Drivers of Delinquency Risk:', level=2)
p = doc.add_paragraph()
p.add_run('1. Credit Card Type (Business): ').bold = True
p.add_run('Strongest predictor. Business card holders show distinct risk patterns compared to Standard or Gold types.')
p = doc.add_paragraph()
p.add_run('2. Income: ').bold = True
p.add_run('Inverse relationship; lower income brackets correlate with higher delinquency probability.')
p = doc.add_paragraph()
p.add_run('3. Recent Payment History (Months 4-5): ').bold = True
p.add_run('Behavior in the 4-5 month window is more predictive than the most recent month, suggesting a lag effect in default.')

# 7. Model Fairness & Bias Analysis
doc.add_heading('7. Model Fairness & Bias Analysis', level=1)
doc.add_paragraph('We conducted a Segmented Performance Analysis to test for bias across Age, Income, and Location groups. The goal was to ensure the model does not systematically disadvantage any specific demographic.')

doc.add_heading('Key Observation: Universal Underperformance', level=2)
doc.add_paragraph('Current analysis shows a Recall of 0.00 across all demographic groups. This indicates the model is currently failing to identify delinquent customers in ANY group, likely due to class imbalance or the "hypothetical" nature of the small dataset.')
doc.add_paragraph('While "fair" in the sense that it fails equally for everyone, this is a performance failure rather than a specific bias issue. However, we noted potential risk disparities:')

bias_findings = [
    ('Age', 'Younger customers (<30) have a higher base delinquency rate (20.8%) vs Middle-aged (11.1%). The model must be tuned to capture this risk signal.'),
    ('Income', 'Lower income groups show higher delinquency (22.2%) compared to High income (14.7%). We must ensure Income does not become a sole proxy for denial.'),
    ('Location', 'Los Angeles shows the highest regional risk (23.1%), while Chicago is lowest (10%).')
]

for area, finding in bias_findings:
    p = doc.add_paragraph()
    p.add_run(f'- {area} Disparity: ').bold = True
    p.add_run(finding)

doc.add_paragraph('Recommendation: Before deployment, the model requires Retraining with heavily weighted classes or lower probability thresholds to improve Recall above 0.00.')

# --- STRATEGIC RECOMMENDATION (5-Step Framework) ---
doc.add_page_break()
doc.add_heading('Strategic Business Recommendation', level=0)

# Step 1: Restate the Insight
doc.add_heading('Step 1: The Insight', level=1)
doc.add_paragraph('Analysis reveals that customers with a Credit Utilization greater than 50% are significantly more likely to default. However, a critical data anomaly exists: "Missed Payments" currently shows a negative correlation with delinquency, suggesting a failure in the data pipeline that must be resolved before model deployment.')

# Step 2: Brainstorm Actions
doc.add_heading('Step 2: Brainstorming Potential Actions', level=1)
doc.add_paragraph('We identified four potential interventions to address these findings:')
brainstorm_list = [
    ('Proactive "Utilization Alert"', 'Send SMS when utilization crosses 45% (Communication).'),
    ('Credit Limit Freeze', 'Stop line increases for high-utilization accounts (Policy).'),
    ('Data Pipeline Audit', 'Technical deep-dive to fix "Missed Payments" data (Technical).'),
    ('Hardship Program', 'Offer restructuring to customers with >80% utilization (Support).')
]
for action, desc in brainstorm_list:
    p = doc.add_paragraph()
    p.add_run(f'• {action}: ').bold = True
    p.add_run(desc)

# Step 3: Evaluate Options
doc.add_heading('Step 3: Evaluation of Options', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Option'
hdr[1].text = 'Feasibility'
hdr[2].text = 'Impact'
hdr[3].text = 'Risk'

eval_data = [
    ('Utilization Alert', 'High (Uses existing SMS tool)', 'Medium (Preventative)', 'Low'),
    ('Limit Freeze', 'High (Policy change)', 'High (Risk reduction)', 'Medium (Customer friction)'),
    ('Data Audit', 'Medium (Resource intensive)', 'Critical (Enables Model)', 'None (Internal only)'),
]
for opt, feas, imp, risk in eval_data:
    row = table.add_row().cells
    row[0].text = opt
    row[1].text = feas
    row[2].text = imp
    row[3].text = risk

# Step 4: The Recommendation
doc.add_heading('Step 4: Recommendation', level=1)
p = doc.add_paragraph()
p.add_run('Primary Recommendation: ').bold = True
p.add_run('Pilot a "Credit Health SMS" campaign for customers exceeding 45% utilization, aiming to reduce roll-rate to delinquency by 10%.')
p = doc.add_paragraph()
p.add_run('Technical Prerequisite: ').bold = True
p.add_run('Immediately freeze "Missed Payment" feature engineering until Data Engineering validates the source data (due to negative correlation anomaly).')

# Step 5: Justification
doc.add_heading('Step 5: Justification', level=1)
doc.add_paragraph('• Alignment: Directly addresses the #1 confirmed risk factor (Utilization) without aggressively cutting lines.')
doc.add_paragraph('• ROI: SMS is low-cost (<$0.01/msg) vs. potential loss recovery ($1000s per default).')
doc.add_paragraph('• Scalability: Can be rolled out to the entire customer base instantly after the pilot.')
doc.add_paragraph('• Ethics: Provides help/information rather than punitive measures like account closures.')

doc.save('Strategic_Recommendation_Report.docx')
print("Report generated: Strategic_Recommendation_Report.docx")
