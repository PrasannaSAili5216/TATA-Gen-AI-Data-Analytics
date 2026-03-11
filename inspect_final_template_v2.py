
# 
from docx import Document
import sys

# Force UTF-8 output wrapper
class Unbuffered(object):
   def __init__(self, stream):
       self.stream = stream
   def write(self, data):
       self.stream.write(data)
       self.stream.flush()
   def __getattr__(self, attr):
       return getattr(self.stream, attr)

sys.stdout = Unbuffered(open('template_structure.txt', 'w', encoding='utf-8'))

try:
    doc = Document('TASK_DETAILS/TASK_3/Updated_Business_Summary_Report_Template.docx')
    print("--- Template Structure ---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"P{i}: {text}")
    
    print("\n--- Tables ---")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}:")
        for j, row in enumerate(table.rows):
            # Join cell text with pipe for clearer table structure view
            row_text = " | ".join([cell.text.strip().replace('\n', ' ') for cell in row.cells])
            print(f"  R{j}: {row_text}")
            
except Exception as e:
    print(f"Error reading template: {e}")
